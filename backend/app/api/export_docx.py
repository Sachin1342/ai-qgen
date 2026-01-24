from fastapi import APIRouter
from pydantic import BaseModel
from docx import Document
from pathlib import Path
import uuid

router = APIRouter()
OUT = Path("outputs/docx")
OUT.mkdir(parents=True, exist_ok=True)

class Export(BaseModel):
    title: str
    questions: list

@router.post("/export/docx")
def export_docx(req: Export):
    name = f"{uuid.uuid4()}.docx"
    doc = Document()
    doc.add_heading(req.title)

    for i, q in enumerate(req.questions, 1):
        doc.add_paragraph(f"{i}. {q}")

    doc.save(OUT / name)
    return {"file": name}
