"""
Export API endpoints - exports questions to PDF and DOCX formats
"""
import os
import uuid
from flask import Blueprint, request, jsonify, send_file
from pathlib import Path

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
except ImportError:
    reportlab = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

export_bp = Blueprint('export', __name__, url_prefix='/api/export')

# Configuration
OUTPUT_FOLDER = 'backend/outputs'
PDF_FOLDER = os.path.join(OUTPUT_FOLDER, 'pdf')
DOCX_FOLDER = os.path.join(OUTPUT_FOLDER, 'docx')

# Ensure output folders exist
Path(PDF_FOLDER).mkdir(parents=True, exist_ok=True)
Path(DOCX_FOLDER).mkdir(parents=True, exist_ok=True)


@export_bp.route('/pdf', methods=['POST'])
def export_pdf():
    """
    Export questions to PDF format
    """
    try:
        data = request.get_json()
        title = data.get('title', 'Generated Questions')
        questions = data.get('questions', [])

        if not questions:
            return jsonify({'detail': 'No questions to export'}), 400

        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_path = os.path.join(PDF_FOLDER, f'{file_id}.pdf')

        # Create simple PDF using reportlab or fallback to text
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch

            doc = SimpleDocTemplate(file_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()

            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='#1a1a1a',
                spaceAfter=30,
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3 * inch))

            # Add questions
            for i, q in enumerate(questions, 1):
                question_text = f"<b>{i}. {q.get('question', '')}</b>"
                story.append(Paragraph(question_text, styles['Normal']))

                if q.get('options') and len(q['options']) > 0:
                    for j, option in enumerate(q['options']):
                        opt_letter = chr(97 + j)  # a, b, c, d...
                        story.append(Paragraph(f"&nbsp;&nbsp;{opt_letter}) {option}", styles['Normal']))
                    if q.get('answer'):
                        story.append(Paragraph(f"<b>Answer:</b> {q['answer']}", styles['Normal']))
                else:
                    story.append(Paragraph("[Descriptive answer required]", styles['Normal']))

                story.append(Spacer(1, 0.2 * inch))

            doc.build(story)

        except ImportError:
            # Fallback to simple text file if reportlab not available
            with open(file_path.replace('.pdf', '.txt'), 'w') as f:
                f.write(f"{title}\n")
                f.write("=" * 50 + "\n\n")
                for i, q in enumerate(questions, 1):
                    f.write(f"{i}. {q.get('question', '')}\n")
                    if q.get('options'):
                        for j, option in enumerate(q['options']):
                            f.write(f"   {chr(97 + j)}) {option}\n")
                    f.write("\n")

        return jsonify({
            'download_url': f'/download/pdf/{file_id}',
            'file_id': file_id,
            'message': 'PDF exported successfully'
        }), 200

    except Exception as e:
        return jsonify({'detail': f'PDF export error: {str(e)}'}), 500


@export_bp.route('/docx', methods=['POST'])
def export_docx():
    """
    Export questions to DOCX format
    """
    try:
        data = request.get_json()
        title = data.get('title', 'Generated Questions')
        questions = data.get('questions', [])

        if not questions:
            return jsonify({'detail': 'No questions to export'}), 400

        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_path = os.path.join(DOCX_FOLDER, f'{file_id}.docx')

        # Create DOCX document
        if Document:
            doc = Document()

            # Add title
            title_heading = doc.add_heading(title, 0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add questions
            for i, q in enumerate(questions, 1):
                # Add question
                question_para = doc.add_paragraph(f"{i}. {q.get('question', '')}", style='Heading 3')

                # Add options if MCQ
                if q.get('options') and len(q['options']) > 0:
                    for j, option in enumerate(q['options']):
                        opt_letter = chr(97 + j)
                        doc.add_paragraph(f"{opt_letter}) {option}", style='List Bullet')

                    # Add answer
                    if q.get('answer'):
                        answer_para = doc.add_paragraph(f"Answer: {q['answer']}")
                        answer_para.runs[0].bold = True
                else:
                    doc.add_paragraph("[Descriptive answer required]")

                # Add spacing
                doc.add_paragraph()

            doc.save(file_path)
        else:
            raise ImportError("python-docx not installed")

        return jsonify({
            'download_url': f'/download/docx/{file_id}',
            'file_id': file_id,
            'message': 'DOCX exported successfully'
        }), 200

    except Exception as e:
        return jsonify({'detail': f'DOCX export error: {str(e)}'}), 500
